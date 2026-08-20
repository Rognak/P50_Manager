"""Демо-кандидаты в разных стадиях воронки.

Идемпотентно: удаляет существующих кандидатов админа и пересоздаёт.
Для части генерирует DOCX-резюме программно (python-docx) и подкладывает
заглушки AI-сводок/feedback (без реального вызова AI).

Запуск:
    uv run python -m scripts.seed_candidates
"""

import asyncio
import io
import sys
from datetime import UTC, date, datetime, timedelta

from docx import Document
from sqlalchemy import delete, select

from app.db import SessionLocal
from app.models.candidate import CandidateProfile
from app.models.employee import Employee
from app.models.mpk import Grade, Role
from app.models.user import User
from app.self_review.docx_render import extract_docx_text

ADMIN_EMAIL = "admin@example.com"
DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def make_resume_docx(
    full_name: str,
    position: str,
    years: int,
    skills: list[str],
    jobs: list[tuple[str, str, list[str]]],
) -> bytes:
    """Генерирует простой DOCX-резюме для тестов."""
    doc = Document()
    doc.add_heading(full_name, level=0)
    doc.add_paragraph(f"Желаемая позиция: {position}")
    doc.add_paragraph(f"Опыт работы: {years} лет")

    doc.add_heading("Ключевые навыки", level=1)
    for s in skills:
        doc.add_paragraph(f"• {s}", style="List Bullet")

    doc.add_heading("Опыт работы", level=1)
    for company, period, achievements in jobs:
        doc.add_heading(f"{company} — {period}", level=2)
        for a in achievements:
            doc.add_paragraph(f"• {a}", style="List Bullet")

    doc.add_heading("Образование", level=1)
    doc.add_paragraph("МГТУ им. Баумана, ФН, бакалавр прикладной математики")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- спецификации кандидатов ----------

CANDIDATES: list[dict] = [
    {
        "full_name": "Демо Кандидат 01",
        "email": "ivanova.maria@candidate.demo",
        "position": "Frontend Senior",
        "source": "LinkedIn",
        "expected_role": "Frontend web разработчик",
        "expected_grade": "Senior",
        "stage": "interview",
        "feedback_decision": None,
        "resume": {
            "years": 7,
            "skills": ["React", "TypeScript", "Next.js", "GraphQL", "Tailwind", "Webpack/Vite"],
            "jobs": [
                (
                    "Яндекс — 2020–2024",
                    "2020–2024",
                    [
                        "Tech Lead команды из 6 фронтов",
                        "Архитектура нового SSR-стека на Next.js — снизили TTFB в 2 раза",
                        "Внедрил design-system с Storybook, переиспользуют все продукты",
                    ],
                ),
                (
                    "Avito — 2018–2020",
                    "2018–2020",
                    [
                        "Разработка маркетплейса B2B, 5М MAU",
                        "Перевёл legacy-jQuery на React, +30% к скорости разработки",
                    ],
                ),
            ],
        },
        "ai_summary": (
            "## Краткое резюме\n"
            "Кандидат — Senior Frontend с 7-летним опытом, специализация на React/TypeScript "
            "и архитектурных решениях. Был tech lead'ом в Яндексе.\n\n"
            "## Ключевые навыки\n"
            "**Языки/фреймворки**: TypeScript, React, Next.js, GraphQL.\n"
            "**Tooling**: Webpack, Vite, Storybook, Tailwind.\n"
            "**Soft**: лидерство (вёл 6 человек), архитектурные решения, дизайн-системы.\n\n"
            "## Опыт работы\n"
            "- Яндекс (2020–2024): Tech Lead, SSR на Next.js (-50% TTFB), внедрение DS\n"
            "- Avito (2018–2020): миграция jQuery→React, B2B маркетплейс\n\n"
            "## Образование и сертификации\n"
            "МГТУ им. Баумана, бакалавр прикладной математики.\n\n"
            "## Что уточнить на интервью\n"
            "1. Опыт менторства: сколько разработчиков довела до повышения?\n"
            "2. Конкретика по «-50% TTFB» — что именно делали?\n"
            "3. Интерес к нашему стеку (Vite + Tailwind + React)?\n"
            "4. Готовность к лидерству или хочет «руками покодить»?\n"
        ),
    },
    {
        "full_name": "Демо Кандидат 02",
        "email": "petrov.d@candidate.demo",
        "position": "Backend Python Middle",
        "source": "Реферал",
        "expected_role": "Backend Python разработчик",
        "expected_grade": "Middle",
        "stage": "hired",
        "feedback_decision": "positive",
        "resume": {
            "years": 4,
            "skills": ["Python", "FastAPI", "PostgreSQL", "Redis", "Docker", "Kubernetes"],
            "jobs": [
                (
                    "Тинькофф — 2021–2024",
                    "2021–2024",
                    [
                        "Разработка платёжных микросервисов на FastAPI",
                        "Миграция монолита (Django) на async — рост RPS в 3 раза",
                        "Менторство 2-х джунов",
                    ],
                ),
            ],
        },
        "ai_summary": (
            "## Краткое резюме\n"
            "Backend Middle, 4 года, Python/FastAPI, опыт миграции Django→async.\n\n"
            "## Ключевые навыки\n"
            "Python, FastAPI, PostgreSQL, Redis, Docker, K8s.\n\n"
            "## Опыт работы\n"
            "- Тинькофф (2021–2024): платёжные сервисы, миграция на async, менторство\n\n"
            "## Что уточнить на интервью\n"
            "1. Опыт работы с очередями (RabbitMQ/Kafka)?\n"
            "2. Самые сложные баги в проде, как дебажил?\n"
        ),
        "ai_feedback": (
            "## Решение\n"
            "Рекомендуется к найму. Уверенный Middle с потенциалом роста до Senior за год-полтора.\n\n"
            "## Сильные стороны\n"
            "- Чисто пишет на FastAPI, понимает async-модель.\n"
            "- Внятно объясняет архитектурные решения, было видно опыт миграции монолита.\n"
            "- Доброжелателен, конструктивно принимает фидбек.\n\n"
            "## Слабые места / риски\n"
            "- Не работал с Kafka/RabbitMQ — на старте может быть просадка.\n"
            "- Опыт K8s ограничен (использовал чужие манифесты).\n\n"
            "## Обратная связь для кандидата\n"
            "Спасибо за интервью! Получил оффер. На старте предлагаем погружение в "
            "очереди (Kafka) и в наш K8s-стек — есть план онбординга, наставник Иванов.\n\n"
            "## План онбординга\n"
            "1. Неделя 1: погружение в монорепо, парное программирование с наставником\n"
            "2. Неделя 2-3: первый PR в платёжный сервис\n"
            "3. Месяц 1: Kafka-курс + интеграция первого consumer'а\n"
            "4. Месяц 1: ревью МПК-профиля, согласование плана развития\n"
        ),
    },
    {
        "full_name": "Демо Кандидат 03",
        "email": "sidorova.a@candidate.demo",
        "position": "DevOps инженер",
        "source": "hh.ru",
        "expected_role": "DevOps инженер",
        "expected_grade": "Middle",
        "stage": "rejected",
        "feedback_decision": "negative",
        "rejection_reason": (
            "Сильно расходится по уровню: позиционирует себя как Middle, но "
            "знание K8s ограничено `kubectl apply` и базовыми Helm-чартами. "
            "Опыт CI/CD только с GitLab CI, наш Jenkins не знает. По сетям — "
            "поверхностно. Возможен возврат через 6-9 месяцев."
        ),
        "resume": {
            "years": 3,
            "skills": ["Linux", "Docker", "Kubernetes (basic)", "GitLab CI", "Terraform"],
            "jobs": [
                (
                    "Lamoda — 2022–2024",
                    "2022–2024",
                    [
                        "Поддержка stage/prod на GitLab CI",
                        "Миграция деплоя с manual в k8s",
                    ],
                ),
            ],
        },
        "ai_summary": (
            "## Краткое резюме\n"
            "DevOps Junior+/Middle (по тексту 3 года), GitLab CI, K8s начального уровня.\n\n"
            "## Что уточнить на интервью\n"
            "1. Глубина Kubernetes — что писала сама?\n"
            "2. Опыт с Helm, Argo CD?\n"
            "3. Сетевая инфраструктура: VPC/security groups/iptables?\n"
        ),
    },
    {
        "full_name": "Демо Кандидат 04",
        "email": "kuznetsov.a@candidate.demo",
        "position": "Mobile iOS разработчик",
        "source": "LinkedIn",
        "expected_role": "Mobile iOS разработчик",
        "expected_grade": "Middle+",
        "stage": "screening",
        "feedback_decision": None,
        "resume": {
            "years": 5,
            "skills": [
                "Swift 5+",
                "SwiftUI / UIKit",
                "Combine",
                "Core Data, Realm",
                "MVVM, Clean Architecture",
                "Fastlane, CI/CD (Bitrise)",
            ],
            "jobs": [
                (
                    "СберМаркет — 2021–2024",
                    "2021–2024",
                    [
                        "Разработка фичей покупательского iOS-приложения (5М+ MAU)",
                        "Перевод части экранов с UIKit на SwiftUI, выработка стандартов команды",
                        "Снижение крэш-rate с 0.4% до 0.08% через структурный рефакторинг сетевого слоя",
                    ],
                ),
                (
                    "Кошелёк — 2019–2021",
                    "2019–2021",
                    [
                        "iOS-разработка модулей лояльности и оплаты",
                        "Внедрил Fastlane для автоматизации релизного цикла (TestFlight + App Store)",
                    ],
                ),
            ],
        },
    },
    {
        "full_name": "Демо Кандидат 05",
        "email": "novikova.e@candidate.demo",
        "position": "Тестировщик (авто)",
        "source": "Реферал",
        "expected_role": "Тестировщик",
        "expected_grade": "Middle",
        "stage": "new",
        "feedback_decision": None,
        "resume": {
            "years": 4,
            "skills": [
                "Python (pytest, requests)",
                "REST API testing",
                "Postman, Insomnia",
                "Selenium / Playwright",
                "Allure-отчёты",
                "PostgreSQL (для проверок)",
                "Charles, Wireshark (basic)",
            ],
            "jobs": [
                (
                    "Альфа-Банк — 2022–2024",
                    "2022–2024",
                    [
                        "Авто-тестирование REST API мобильного банка (~200 тест-кейсов)",
                        "Внедрила интеграцию Allure в Jenkins-пайплайн команды",
                        "Сократила регресс с 3 дней до 4 часов за счёт авто-тестов",
                    ],
                ),
                (
                    "Lamoda Tech — 2020–2022",
                    "2020–2022",
                    [
                        "Ручное и авто-тестирование e-commerce, фокус на checkout",
                        "Подготовка и поддержка тест-данных в стейдж-окружениях",
                    ],
                ),
            ],
        },
    },
    {
        "full_name": "Демо Кандидат 06",
        "email": "morozov.i@candidate.demo",
        "position": "Backend Java разработчик",
        "source": "хантеры",
        "expected_role": "Backend Java разработчик",
        "expected_grade": "Senior",
        "stage": "interview",
        "feedback_decision": None,
        "resume": {
            "years": 8,
            "skills": ["Java 17/21", "Spring Boot", "Kafka", "PostgreSQL", "Kubernetes"],
            "jobs": [
                (
                    "Сбер — 2018–2024",
                    "2018–2024",
                    [
                        "Разработка high-load платёжных сервисов",
                        "Tech Lead команды из 5 человек на проекте «Антифрод»",
                        "Миграция с Java 8 на 17, сокращение latency на 40%",
                    ],
                ),
                (
                    "РЖД — 2016–2018",
                    "2016–2018",
                    [
                        "Разработка систем учёта грузоперевозок",
                    ],
                ),
            ],
        },
    },
    {
        "full_name": "Демо Кандидат 07",
        "email": "vasilieva.o@candidate.demo",
        "position": "Системный аналитик",
        "source": "Реферал",
        "expected_role": "Системный аналитик",
        "expected_grade": "Senior",
        "stage": "offer",
        "feedback_decision": "positive",
        "resume": {
            "years": 6,
            "skills": ["BPMN", "UML", "SQL", "Confluence", "Jira", "ArchiMate"],
            "jobs": [
                (
                    "X5 Retail — 2020–2024",
                    "2020–2024",
                    [
                        "Анализ требований для омниканальной платформы",
                        "Постановка задач командам разработки (8 продуктов)",
                    ],
                ),
            ],
        },
        "ai_feedback": (
            "## Решение\n"
            "Рекомендуется к найму. Senior-уровень подтверждается опытом в X5.\n\n"
            "## Сильные стороны\n"
            "- Чёткое мышление, BPMN/UML на автомате.\n"
            "- Опыт с большой командой и омниканальной архитектурой.\n\n"
            "## Слабые места / риски\n"
            "- Ожидания по зарплате выше нашей вилки на ~10%.\n\n"
            "## Обратная связь для кандидата\n"
            "Готовим оффер с обсуждением вилки — обсудим в течение недели.\n"
        ),
    },
    {
        "full_name": "Демо Кандидат 08",
        "email": "tikhonov.a@candidate.demo",
        "position": "Backend Python Junior",
        "source": "hh.ru",
        "expected_role": "Backend Python разработчик",
        "expected_grade": "Junior",
        "stage": "rejected",
        "feedback_decision": "negative",
        "rejection_reason": (
            "Базовые знания Python есть, но опыт ограничен учебными проектами. "
            "На задаче на FastAPI не справился — путался в async/await, не понимал "
            "ORM-сессии. Рекомендую вернуться через 3-6 месяцев после стажировки."
        ),
        "resume": {
            "years": 1,
            "skills": ["Python", "Flask (учебно)", "PostgreSQL"],
            "jobs": [
                (
                    "Самозанятость / pet-проекты — 2023–2024",
                    "2023–2024",
                    [
                        "Telegram-бот с обработкой платежей (учебный)",
                    ],
                ),
            ],
        },
    },
]


async def main() -> None:
    async with SessionLocal() as session:
        admin = (
            await session.execute(select(User).where(User.email == ADMIN_EMAIL))
        ).scalar_one_or_none()
        if admin is None:
            print(f"!! не найден {ADMIN_EMAIL} — сначала seed-admin", file=sys.stderr)
            sys.exit(1)

        # очистка кандидатов админа
        existing_ids_q = await session.execute(
            select(Employee.id).where(Employee.owner_id == admin.id, Employee.kind == "candidate")
        )
        existing_ids = [eid for (eid,) in existing_ids_q.all()]
        if existing_ids:
            await session.execute(
                delete(CandidateProfile).where(CandidateProfile.employee_id.in_(existing_ids))
            )
            await session.execute(delete(Employee).where(Employee.id.in_(existing_ids)))
            await session.commit()

        # справочники
        roles = {r.name: r.id for r in (await session.execute(select(Role))).scalars()}
        grades = {g.code: g.id for g in (await session.execute(select(Grade))).scalars()}

        now = datetime.now(UTC)
        today = date.today()

        for spec in CANDIDATES:
            emp = Employee(
                full_name=spec["full_name"],
                email=spec["email"],
                position=spec["position"],
                owner_id=admin.id,
                kind="candidate",
            )
            if spec["stage"] == "hired":
                emp.kind = "employee"
                emp.hired_at = today - timedelta(days=15)
                # роли/грейд переносятся на сотрудника
                if spec.get("expected_role") in roles:
                    emp.role_id = roles[spec["expected_role"]]
                if spec.get("expected_grade") in grades:
                    emp.grade_id = grades[spec["expected_grade"]]
            session.add(emp)
            await session.flush()

            prof = CandidateProfile(
                employee_id=emp.id,
                stage=spec["stage"],
                source=spec.get("source"),
                expected_role_id=roles.get(spec.get("expected_role") or ""),
                expected_grade_id=grades.get(spec.get("expected_grade") or ""),
                feedback_decision=spec.get("feedback_decision"),
                rejection_reason_md=spec.get("rejection_reason"),
                ai_resume_summary_md=spec.get("ai_summary"),
                ai_feedback_md=spec.get("ai_feedback"),
            )

            # резюме
            res = spec.get("resume")
            if res is not None:
                blob = make_resume_docx(
                    spec["full_name"],
                    spec["position"],
                    res["years"],
                    res["skills"],
                    res["jobs"],
                )
                prof.resume_data = blob
                prof.resume_filename = f"{spec['full_name'].split()[0]}_резюме.docx"
                prof.resume_content_type = DOCX_TYPE
                prof.resume_size_bytes = len(blob)
                prof.resume_uploaded_at = now - timedelta(days=3)
                prof.resume_text = extract_docx_text(blob)

            session.add(prof)

        await session.commit()

    print(f"Загружено кандидатов: {len(CANDIDATES)}")
    by_stage: dict[str, int] = {}
    for c in CANDIDATES:
        by_stage[c["stage"]] = by_stage.get(c["stage"], 0) + 1
    for stage, n in by_stage.items():
        print(f"  {stage:12s}: {n}")
    have_resume = sum(1 for c in CANDIDATES if c.get("resume"))
    have_summary = sum(1 for c in CANDIDATES if c.get("ai_summary"))
    have_feedback = sum(1 for c in CANDIDATES if c.get("ai_feedback"))
    print(f"  с резюме:    {have_resume}")
    print(f"  с AI-сводкой: {have_summary}")
    print(f"  с AI-feedback: {have_feedback}")


if __name__ == "__main__":
    asyncio.run(main())
