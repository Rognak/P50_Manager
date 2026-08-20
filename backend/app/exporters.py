"""Экспорт markdown-документов в DOCX и print-ready HTML (для сохранения как PDF из браузера)."""

import io
import re
from html import escape

import markdown as md_lib
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


_CODE_FONT = "Menlo"
_BODY_FONT = "Calibri"


def _set_base_style(doc: DocumentObject) -> None:
    style = doc.styles["Normal"]
    style.font.name = _BODY_FONT
    style.font.size = Pt(11)


def _add_inline(paragraph, text: str) -> None:
    """Очень простой парсер инлайн-разметки: **bold**, *italic*, `code`, [text](url)."""
    # Порядок: сначала код, потом ссылки, потом жирный/курсив
    pattern = re.compile(
        r"(\*\*[^*\n]+?\*\*)"  # **bold**
        r"|(\*[^*\n]+?\*)"  # *italic*
        r"|(`[^`\n]+?`)"  # `code`
        r"|(\[[^\]]+?\]\([^)]+?\))"  # [text](url)
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = _CODE_FONT
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif token.startswith("["):
            link_m = re.match(r"\[([^\]]+?)\]\(([^)]+?)\)", token)
            if link_m:
                run = paragraph.add_run(link_m.group(1))
                run.font.color.rgb = RGBColor(0x10, 0x4E, 0x8B)
                run.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx(md_text: str, title: str) -> bytes:
    """Построчный парсинг markdown в docx. Поддерживает заголовки, списки,
    code блоки ```, жирный/курсив/inline code/ссылки."""
    doc = Document()
    _set_base_style(doc)

    # Заголовок документа
    h = doc.add_heading(title, level=0)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    lines = md_text.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            run = p.add_run("\n".join(code_lines))
            run.font.name = _CODE_FONT
            run.font.size = Pt(10)
            if lang:
                lang_p = doc.add_paragraph()
                lang_run = lang_p.add_run(f"[{lang}]")
                lang_run.font.size = Pt(8)
                lang_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            continue

        # Пустая строка
        if not stripped:
            i += 1
            continue

        # Заголовки
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        # Списки: - / * / 1.
        elif re.match(r"^(-|\*)\s+", stripped):
            content = re.sub(r"^(-|\*)\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, content)
        elif re.match(r"^\d+\.\s+", stripped):
            content = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, content)
        else:
            p = doc.add_paragraph()
            _add_inline(p, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


_PRINT_CSS = """
body {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
  color: #222;
  max-width: 760px;
  margin: 40px auto;
  padding: 0 24px;
  line-height: 1.55;
}
h1 { font-size: 26pt; margin: 0 0 16pt; }
h2 { font-size: 16pt; margin: 20pt 0 8pt; border-bottom: 1px solid #ddd; padding-bottom: 4pt; }
h3 { font-size: 13pt; margin: 14pt 0 6pt; }
p, li { font-size: 11pt; }
code { background: #f1f1f2; padding: 1px 6px; border-radius: 3px; font-size: 10.5pt; }
pre { background: #f1f1f2; padding: 12px 14px; border-radius: 6px; overflow-x: auto; }
pre code { background: none; padding: 0; }
ul, ol { padding-left: 24px; }
a { color: #104e8b; }
.meta { color: #888; font-size: 10pt; margin-bottom: 20pt; }
@media print {
  body { margin: 0; }
}
"""

_AUTO_PRINT_JS = """
window.addEventListener('load', () => {
  setTimeout(() => window.print(), 250);
});
"""


def markdown_to_print_html(md_text: str, title: str, meta: str | None = None) -> str:
    """Генерирует полноценный HTML с print-ready CSS и auto-print JS.
    Пользователь открывает URL в браузере → автоматом появляется диалог печати → «Сохранить как PDF»."""
    body = md_lib.markdown(
        md_text,
        extensions=["extra", "sane_lists", "codehilite", "fenced_code"],
    )
    meta_html = f'<div class="meta">{escape(meta)}</div>' if meta else ""
    return f"""<!doctype html>
<html lang="ru">
<head>
<meta charset="utf-8">
<title>{escape(title)}</title>
<style>{_PRINT_CSS}</style>
</head>
<body>
<h1>{escape(title)}</h1>
{meta_html}
{body}
<script>{_AUTO_PRINT_JS}</script>
</body>
</html>"""
